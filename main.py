import time
from PIL import ImageGrab
import requests
import selenium
from selenium import webdriver
import secert
from selenium.webdriver.common.action_chains import ActionChains
from selenium.webdriver.common.keys import Keys
from docx import Document
from pykeyboard import PyKeyboard

k = PyKeyboard()

target = "https://moodle.scnu.edu.cn/mod/quiz/review.php?attempt=1244357&cmid=83774"
words = "1260004"

docu = Document(r"D:\Chorme\PyGaoShu\1.docx")

dr = webdriver.Chrome()
# dr.get("https://moodle.scnu.edu.cn/mod/quiz/review.php?attempt=1203635&cmid=81634")
dr.get(target)

# dr.find_element_by_id("question-1219243-1")
# print(dr.find_element_by_id("question-1219243-1").text)
dr.find_element_by_id("ssobtn").click()
dr.find_element_by_id("account").clear()
dr.find_element_by_id("account").send_keys(secert.name)
dr.find_element_by_id("password").clear()
dr.find_element_by_id("password").send_keys(secert.passwordID)
dr.find_element_by_id("rancode").clear()
randcodes = input("input random code:")

dr.find_element_by_id("rancode").send_keys(randcodes)
dr.find_element_by_class_name("login").click()
time.sleep(2)
dr.find_element_by_id("gotoAppBtn").click()


# print(dr.find_element_by_id("question-1219243-1").text)

def saveimg(url, x, flag):
    while True:
        try:
            action = ActionChains(dr).move_to_element(url)
            action.context_click(url)
            # action.send_keys(Keys.ARROW_DOWN)
            action.perform()
            print("mouse right press")
            # action.send_keys('Y')
            # action.perform()
            time.sleep(0.2)
            k.press_key('y')
            k.release_key('y')
            time.sleep(0.2)
            print("key press y")
            image = ImageGrab.grabclipboard()
            # p1 = cv2.resize(image, (int(a[1] / 1.5), int(a[0] / 1.5)),interpolation=cv2.INTER_CUBIC)
            sizee = image.size
            image = image.resize((int(sizee[0] * 0.5), int(sizee[1] * 0.5)))
            image.save('2.png')
            print("save img in 2.png")
            time.sleep(0.5)
            break
        except:
            print("Error")
            continue

    if flag == 1:
        docu.add_paragraph("第" + str(x) + "题")
        # run = docu.tables[0].cell(0, 0).paragraphs[0].add_run()
    else:
        docu.add_paragraph("正确答案为：")
    docu.add_picture("2.png")
    docu.save("1.docx")
    print("save in 1.docx")


def savetext(url, x, flag):
    if flag == 1:
        docu.add_paragraph("第" + str(x) + "题")
    else:
        print()
        # docu.add_paragraph("正确答案为：")
    docu.add_paragraph(url.text)
    docu.save("1.docx")


def downloadpic(s, x):
    try:
        url = dr.find_element_by_xpath("//div[@id='question-" + s + "-" + str(x) +
                                       "']"
                                       "/div[@class='content']"
                                       "/div[@class='formulation clearfix']"
                                       "/div[@class='qtext']/p/img")
        print("题目为图片")
        flag = True
    except:
        url = dr.find_element_by_xpath("//div[@id='question-" + s + "-" + str(x) +
                                       "']"
                                       "/div[@class='content']"
                                       "/div[@class='formulation clearfix']"
                                       "/div[@class='qtext']/p")
        print("题目为文字")
        flag = False

    if flag:
        saveimg(url, x, 1)
    else:
        savetext(url, x, 1)


    try:
        ans = dr.find_element_by_xpath("//div[@id='question-" + s + "-" + str(x) +
                                       "']"
                                       "/div[@class='content']"
                                       "/div[@class='outcome clearfix']"
                                       "/div[@class='feedback']"
                                       "/div[@class='rightanswer']/img")
        print("答案为图片")
        flag = True
        # saveimg(ans, x, 2)
    except:
        ans = dr.find_element_by_xpath("//div[@id='question-" + s + "-" + str(x) +
                                       "']"
                                       "/div[@class='content']"
                                       "/div[@class='outcome clearfix']"
                                       "/div[@class='feedback']"
                                       "/div[@class='rightanswer']")
        print("答案为文字")
        flag = False

    if flag:
        saveimg(ans, x, 2)
    else:
        savetext(ans, x, 2)


for i in range(1, 21):
    print(f"正在存储第{i}题")
    downloadpic(words, i)
    print()

dr.close()
# print(url)
# dr.quit()
#
# # 通过requests发送一个get请求到图片地址，返回的响应就是图片内容
# r = requests.get(url)
#
# with open('1.png', 'wb') as f:
#     # 对于图片类型的通过r.content方式访问响应内容，将响应内容写入baidu.png中
#     print(r.content)
#     f.write(r.content)
